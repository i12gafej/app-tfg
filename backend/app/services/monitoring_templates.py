import logging
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

logger = logging.getLogger(__name__)

def add_page_number(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_bookmark(paragraph, bookmark_name):
    
    tag = paragraph._p
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), '0')
    start.set(qn('w:name'), bookmark_name)
    tag.insert(0, start)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), '0')
    tag.append(end)

def add_hyperlink(paragraph, text, bookmark):
    
    part = paragraph.part
    r_id = part.relate_to('#' + bookmark, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    
    b = OxmlElement('w:b')
    b.set(qn('w:val'), 'true')
    rPr.append(b)
    
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def generate_monitoring_template(material_topics):
    logger.info("Iniciando generación de plantilla DOCX bonita y funcional")
    try:
        doc = Document()
        
        for section in doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)
            add_page_number(section)
        
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        
        doc.add_heading('Índice', 0)
        for topic in material_topics:
            
            idx_p = doc.add_paragraph()
            add_hyperlink(idx_p, f"{topic['name']}", f"asunto_{topic['id']}")
            for objective in topic['objectives']:
                for action in objective['actions']:
                    idx_a = doc.add_paragraph(style='List Bullet')
                    add_hyperlink(idx_a, action['description'][:50] + '...', f"accion_{action['id']}")
        doc.add_page_break()

        
        for topic in material_topics:
            
            asunto_p = doc.add_paragraph()
            asunto_p.style = doc.styles['Heading 1']
            asunto_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            asunto_p.add_run(topic['name']).bold = True
            add_bookmark(asunto_p, f"asunto_{topic['id']}")
            doc.add_paragraph()  
            
            asunto_info = [
                ('Prioridad', topic['priority']),
                ('Objetivo principal', topic['main_objective'])
            ]
            for label, value in asunto_info:
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(value)
            doc.add_paragraph()  
            
            for objective in topic['objectives']:
                for action in objective['actions']:
                    
                    action_p = doc.add_paragraph()
                    action_p.style = doc.styles['Heading 2']
                    action_p.add_run(action['description'][:80]).bold = True
                    add_bookmark(action_p, f"accion_{action['id']}")
                    
                    fields = [
                        ('Ejecución', action.get('execution_time', 'No definido')),
                        ('Impacto principal ODS', action['main_impact'] or 'No definido'),
                        ('Indicadores de rendimiento', ', '.join([ind['name'] for ind in action['indicators']])),
                        ('Dificultad', action['difficulty']),
                        ('Responsable equipo', objective['responsible'] or 'No definido'),
                        ('Objetivo específico', objective['description']),
                        ('Resultados', ''),
                        ('Notas', '')
                    ]
                    for label, value in fields:
                        p = doc.add_paragraph()
                        p.add_run(f"{label}: ").bold = True
                        p.add_run(value)
                    
                    now = datetime.datetime.now()
                    year = str(now.year)
                    month = now.month
                    if 1 <= month <= 3:
                        trimestre = "Enero-Marzo"
                    elif 4 <= month <= 6:
                        trimestre = "Abril-Junio"
                    elif 7 <= month <= 9:
                        trimestre = "Julio-Septiembre"
                    else:
                        trimestre = "Octubre-Diciembre"
                    table = doc.add_table(rows=1, cols=2)
                    table.style = 'Light List Accent 1'
                    table.cell(0, 0).text = f'Año: {year}'
                    table.cell(0, 1).text = f'Trimestre: {trimestre}'
                    doc.add_paragraph()  
            doc.add_page_break()
        logger.info("Plantilla DOCX generada correctamente")
        return doc
    except Exception as e:
        logger.error(f"Error al generar plantilla DOCX: {str(e)}", exc_info=True)
        raise
